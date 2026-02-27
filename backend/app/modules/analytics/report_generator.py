"""Enhanced Report Generator Service - Generate professional DOCX and PDF reports with charts."""

import io
import os
from datetime import datetime
from typing import Dict, List, Any, BinaryIO
from io import BytesIO

# Optional imports with graceful fallback
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import matplotlib
    matplotlib.use('Agg')  # Use non-interactive backend
    import matplotlib.pyplot as plt
    import matplotlib.dates as mdates
    from matplotlib.figure import Figure
    import numpy as np
    MATPLOTLIB_AVAILABLE = True
except ImportError:
    MATPLOTLIB_AVAILABLE = False

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak, Image
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    from reportlab.graphics.shapes import Drawing
    from reportlab.graphics.charts.barcharts import VerticalBarChart, HorizontalBarChart
    from reportlab.graphics.charts.linecharts import HorizontalLineChart
    from reportlab.graphics.charts.piecharts import Pie
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

from backend.app.shared.base_service import BaseService


class ReportGenerator(BaseService):
    """Generate professional reports in DOCX and PDF formats with comprehensive charts."""
    
    def __init__(self):
        """Initialize report generator."""
        super().__init__()
        self.logger.info("ReportGenerator initialized")
        self.chart_color_palette = ['#1f4788', '#2874A6', '#5499C7', '#7FB3D5', '#A9CCE3']
    
    def generate_docx_report(self, cost_analysis_data: Dict, filters: Dict) -> BinaryIO:
        """
        Generate a comprehensive DOCX report with cost analysis data and charts.
        
        Args:
            cost_analysis_data: Data from cost analysis API
            filters: Applied filters for the analysis
            
        Returns:
            BytesIO object containing the DOCX file
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for DOCX report generation. "
                "Please install it: pip install python-docx"
            )
        
        doc = Document()
        
        # Add styles
        self._setup_styles(doc)
        
        # Title
        self._add_title(doc, "Pharmaceutical Cost Analysis Report")
        
        # Executive Summary
        self._add_executive_summary(doc, cost_analysis_data, filters)
        
        # Filters Applied
        self._add_filters_section(doc, filters)
        
        # Cost Overview with Chart
        self._add_cost_overview_with_chart(doc, cost_analysis_data)
        
        # Top Cost Drivers with Chart
        self._add_top_cost_drivers_with_chart(doc, cost_analysis_data)
        
        # Cost Trends with Chart
        self._add_cost_trends_with_chart(doc, cost_analysis_data)
        
        # Department Analysis with Chart
        self._add_department_analysis_with_chart(doc, cost_analysis_data)
        
        # Drug Category Analysis with Chart
        self._add_category_analysis_with_chart(doc, cost_analysis_data)
        
        # Bubble Chart Analysis
        self._add_bubble_chart_analysis(doc, cost_analysis_data)
        
        # Detailed Data Tables
        self._add_detailed_sunburst_table(doc, cost_analysis_data)
        
        # Footer
        self._add_footer(doc)
        
        # Save to BytesIO
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output
    
    def generate_pdf_report(self, cost_analysis_data: Dict, filters: Dict) -> BinaryIO:
        """
        Generate a comprehensive PDF report with cost analysis data and charts.
        
        Args:
            cost_analysis_data: Data from cost analysis API
            filters: Applied filters for the analysis
            
        Returns:
            BytesIO object containing the PDF file
        """
        if not REPORTLAB_AVAILABLE:
            raise ImportError(
                "reportlab is required for PDF report generation. "
                "Please install it: pip install reportlab"
            )
        
        output = BytesIO()
        doc = SimpleDocTemplate(output, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
        
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1f4788'),
            spaceAfter=30,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#1f4788'),
            spaceAfter=12,
            spaceBefore=12,
            fontName='Helvetica-Bold'
        )
        
        subheading_style = ParagraphStyle(
            'CustomSubHeading',
            parent=styles['Heading3'],
            fontSize=12,
            textColor=colors.HexColor('#2874A6'),
            spaceAfter=8,
            spaceBefore=8,
            fontName='Helvetica-Bold'
        )
        
        # Title
        story.append(Paragraph("Pharmaceutical Cost Analysis Report", title_style))
        story.append(Spacer(1, 0.2*inch))
        
        # Date and filters info
        report_date = datetime.now().strftime("%B %d, %Y")
        story.append(Paragraph(f"<b>Report Date:</b> {report_date}", styles['Normal']))
        story.append(Spacer(1, 0.1*inch))
        
        # Executive Summary
        story.append(Paragraph("Executive Summary", heading_style))
        summary_text = self._generate_summary_text(cost_analysis_data, filters)
        story.append(Paragraph(summary_text, styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        
        # Filters Applied
        story.append(Paragraph("Analysis Filters", heading_style))
        filters_table = self._create_filters_table(filters)
        story.append(filters_table)
        story.append(Spacer(1, 0.2*inch))
        
        # Cost Overview Table
        story.append(Paragraph("Cost Overview", heading_style))
        overview_table = self._create_cost_overview_table(cost_analysis_data)
        if overview_table:
            story.append(overview_table)
            story.append(Spacer(1, 0.2*inch))
            
            # Cost Overview Chart
            if MATPLOTLIB_AVAILABLE:
                overview_chart = self._create_cost_overview_chart_pdf(cost_analysis_data)
                if overview_chart:
                    story.append(overview_chart)
                    story.append(Spacer(1, 0.2*inch))
        else:
            story.append(Paragraph("No cost overview data available for the selected filters.", styles['Normal']))
            story.append(Spacer(1, 0.2*inch))
        
        story.append(PageBreak())
        
        # Top Cost Drivers Table
        story.append(Paragraph("Top 20 Cost Drivers", heading_style))
        drivers_table = self._create_top_drivers_table(cost_analysis_data)
        if drivers_table:
            story.append(drivers_table)
        story.append(Spacer(1, 0.2*inch))
        
        # Top Cost Drivers Chart
        if MATPLOTLIB_AVAILABLE:
            drivers_chart = self._create_top_drivers_chart_pdf(cost_analysis_data)
            if drivers_chart:
                story.append(drivers_chart)
                story.append(Spacer(1, 0.2*inch))
        
        story.append(PageBreak())
        
        # Cost Trends Section
        story.append(Paragraph("Cost Trends Analysis", heading_style))
        
        # Monthly Trends Table
        story.append(Paragraph("Monthly Cost Trends", subheading_style))
        monthly_table = self._create_monthly_trends_table(cost_analysis_data)
        if monthly_table:
            story.append(monthly_table)
        story.append(Spacer(1, 0.2*inch))
        
        # Cost Trends Chart
        if MATPLOTLIB_AVAILABLE:
            trends_chart = self._create_cost_trends_chart_pdf(cost_analysis_data)
            if trends_chart:
                story.append(trends_chart)
                story.append(Spacer(1, 0.2*inch))
        
        story.append(PageBreak())
        
        # Department Analysis
        story.append(Paragraph("Department Analysis", heading_style))
        dept_table = self._create_department_table(cost_analysis_data)
        if dept_table:
            story.append(dept_table)
        else:
            story.append(Paragraph("No department data available for the selected filters.", styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        
        # Category Analysis
        story.append(Paragraph("Drug Category Analysis", heading_style))
        category_table = self._create_category_table(cost_analysis_data)
        if category_table:
            story.append(category_table)
        else:
            story.append(Paragraph("No category data available for the selected filters.", styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        
        # Bubble Chart Analysis
        story.append(PageBreak())
        story.append(Paragraph("Price vs Quantity Analysis", heading_style))
        bubble_summary = self._generate_bubble_summary(cost_analysis_data)
        story.append(Paragraph(bubble_summary, styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        
        if MATPLOTLIB_AVAILABLE:
            bubble_chart = self._create_bubble_chart_pdf(cost_analysis_data)
            if bubble_chart:
                story.append(bubble_chart)
        
        # Build PDF
        doc.build(story)
        output.seek(0)
        return output
    
    # ========================================================================
    # DOCX Helper Methods with Charts
    # ========================================================================
    
    def _setup_styles(self, doc: Document):
        """Setup document styles."""
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
    
    def _add_title(self, doc: Document, title: str):
        """Add title to document."""
        heading = doc.add_heading(title, 0)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add date
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(f"Report Generated: {datetime.now().strftime('%B %d, %Y at %H:%M')}")
        run.italic = True
        run.font.size = Pt(10)
        
        doc.add_paragraph()  # Spacing
    
    def _add_executive_summary(self, doc: Document, data: Dict, filters: Dict):
        """Add executive summary section."""
        doc.add_heading("Executive Summary", 1)
        
        # Calculate totals
        sunburst = data.get('sunburst', {}).get('data', [])
        drivers = data.get('top_cost_drivers', {}).get('data', [])
        bubble = data.get('bubble_chart', {}).get('data', [])
        
        # Calculate from drug-level items
        drug_items = [item for item in sunburst if item.get('level') == 'drug']
        total_cost = sum(item.get('value', 0) for item in drug_items)
        unique_drugs = len(drug_items)
        unique_depts = len(set(item.get('name') for item in sunburst if item.get('level') == 'department'))
        
        # If sunburst is empty or has no cost, try to use top drivers data
        if (total_cost == 0 or unique_drugs == 0) and drivers:
            total_cost = sum(d.get('total_cost', 0) for d in drivers)
            unique_drugs = len(drivers)
            # Try to get unique departments from drivers
            if unique_depts == 0:
                unique_depts = len(set(d.get('department_id', 'Unknown') for d in drivers if d.get('department_id')))
        
        top_drug = drivers[0].get('drug_name', 'N/A') if drivers else 'N/A'
        top_drug_cost = drivers[0].get('total_cost', 0) if drivers else 0
        
        # Summary text
        p = doc.add_paragraph()
        p.add_run(f"This comprehensive report analyzes pharmaceutical costs from ").font.size = Pt(11)
        run = p.add_run(f"{filters.get('start_date', 'N/A')} to {filters.get('end_date', 'N/A')}")
        run.bold = True
        run.font.size = Pt(11)
        p.add_run(".").font.size = Pt(11)
        
        # Key metrics
        doc.add_paragraph()
        doc.add_heading("Key Metrics", 2)
        
        metrics_data = [
            ["Metric", "Value"],
            ["Total Cost", f"${total_cost:,.2f}"],
            ["Unique Drugs", f"{unique_drugs:,}"],
            ["Departments Analyzed", f"{unique_depts}"],
            ["Top Cost Driver", top_drug],
            ["Top Driver Cost", f"${top_drug_cost:,.2f}"],
        ]
        
        table = doc.add_table(rows=len(metrics_data), cols=2)
        table.style = 'Light Grid Accent 1'
        
        for i, row_data in enumerate(metrics_data):
            row = table.rows[i]
            row.cells[0].text = row_data[0]
            row.cells[1].text = row_data[1]
            if i == 0:  # Header
                for cell in row.cells:
                    cell.paragraphs[0].runs[0].font.bold = True
            else:
                row.cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        doc.add_paragraph()  # Spacing
    
    def _add_filters_section(self, doc: Document, filters: Dict):
        """Add applied filters section."""
        doc.add_heading("Analysis Parameters", 1)
        
        # Handle None values safely
        departments = filters.get('departments') or []
        departments_str = ', '.join(map(str, departments)) if (isinstance(departments, (list, tuple)) and departments) else "All"
        
        drug_categories = filters.get('drug_categories') or []
        categories_str = ', '.join(map(str, drug_categories)) if (isinstance(drug_categories, (list, tuple)) and drug_categories) else "All"
        
        price_min = filters.get('price_min') or 0
        price_max = filters.get('price_max') or 'Unlimited'
        
        rows = [["Parameter", "Value"]]
        rows.append(["Start Date", filters.get('start_date', 'N/A')])
        rows.append(["End Date", filters.get('end_date', 'N/A')])
        rows.append(["Departments", departments_str])
        rows.append(["Price Range", f"${price_min} - ${price_max}"])
        rows.append(["Drug Categories", categories_str])
        
        table = doc.add_table(rows=len(rows), cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Format header
        hdr_cells = table.rows[0].cells
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()  # Spacing
    
    def _add_cost_overview_with_chart(self, doc: Document, data: Dict):
        """Add cost overview section with chart."""
        doc.add_heading("Cost Overview by Department", 1)
        
        sunburst = data.get('sunburst', {}).get('data', [])
        if not sunburst:
            doc.add_paragraph("No data available.")
            return
        
        # Create hierarchy summary
        departments = {}
        for item in sunburst:
            if item.get('level') == 'department':
                dept_name = item.get('name', 'Unknown')
                dept_cost = item.get('value', 0)
                # Count categories under this department
                count = sum(1 for sub in sunburst if sub.get('parent') == dept_name and sub.get('level') == 'category')
                departments[dept_name] = (dept_cost, count)
        
        # Table
        rows = [["Department", "Total Cost", "Categories"]]
        for dept, (cost, count) in sorted(departments.items(), key=lambda x: x[1][0], reverse=True):
            rows.append([dept, f"${cost:,.2f}", str(count)])
        
        table = doc.add_table(rows=len(rows), cols=3)
        table.style = 'Light Grid Accent 1'
        
        # Format header
        hdr_cells = table.rows[0].cells
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].font.bold = True
        
        # Format data cells
        for row in table.rows[1:]:
            row.cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            row.cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        # Add chart if matplotlib is available
        if MATPLOTLIB_AVAILABLE and departments:
            doc.add_paragraph()
            chart_stream = self._create_department_chart(departments)
            if chart_stream:
                doc.add_picture(chart_stream, width=Inches(6))
        
        doc.add_paragraph()  # Spacing
    
    def _add_top_cost_drivers_with_chart(self, doc: Document, data: Dict):
        """Add top cost drivers section with chart."""
        doc.add_heading("Top 20 Cost Drivers", 1)
        
        drivers = data.get('top_cost_drivers', {}).get('data', [])
        if not drivers:
            doc.add_paragraph("No data available.")
            return
        
        # Limit to top 20
        drivers = drivers[:20]
        
        rows = [["Rank", "Drug Name", "Total Cost", "Units", "Avg Price"]]
        for idx, driver in enumerate(drivers, 1):
            rows.append([
                str(idx),
                driver.get('drug_name', 'N/A')[:50],
                f"${driver.get('total_cost', 0):,.2f}",
                str(driver.get('units', 0)),
                f"${driver.get('avg_price', 0):,.2f}"
            ])
        
        table = doc.add_table(rows=len(rows), cols=5)
        table.style = 'Light Grid Accent 1'
        table.autofit = False
        
        # Set column widths
        widths = [Inches(0.6), Inches(2.5), Inches(1.3), Inches(0.9), Inches(1.0)]
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width
        
        # Format header
        hdr_cells = table.rows[0].cells
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].font.bold = True
        
        # Format data cells
        for row in table.rows[1:]:
            row.cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row.cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            row.cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            row.cells[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        # Add chart if matplotlib is available
        if MATPLOTLIB_AVAILABLE:
            doc.add_paragraph()
            chart_stream = self._create_top_drivers_chart(drivers[:10])  # Top 10 for readability
            if chart_stream:
                doc.add_picture(chart_stream, width=Inches(6))
        
        doc.add_paragraph()  # Spacing
    
    def _add_cost_trends_with_chart(self, doc: Document, data: Dict):
        """Add cost trends section with charts."""
        doc.add_heading("Cost Trends Analysis", 1)
        
        cost_trends = data.get('cost_trends', {})
        daily = cost_trends.get('daily', {}).get('data', [])
        monthly = cost_trends.get('monthly', {}).get('data', [])
        
        if not monthly and not daily:
            doc.add_paragraph("No trend data available.")
            return
        
        # Monthly trends table
        doc.add_heading("Monthly Cost Trends", 2)
        rows = [["Period", "Total Cost", "Change %"]]
        prev_cost = None
        for trend in monthly[:12]:  # Last 12 months
            period = trend.get('date', 'N/A')
            cost = trend.get('total_cost', 0)
            
            change = "—"
            if prev_cost is not None and prev_cost > 0:
                pct_change = ((cost - prev_cost) / prev_cost) * 100
                change = f"{pct_change:+.1f}%"
            
            rows.append([period, f"${cost:,.2f}", change])
            prev_cost = cost
        
        table = doc.add_table(rows=len(rows), cols=3)
        table.style = 'Light Grid Accent 1'
        
        # Format header
        hdr_cells = table.rows[0].cells
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].font.bold = True
        
        # Format data cells
        for row in table.rows[1:]:
            row.cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            row.cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        # Add monthly trend chart
        if MATPLOTLIB_AVAILABLE and monthly:
            doc.add_paragraph()
            chart_stream = self._create_trend_chart(monthly, "Monthly Cost Trends")
            if chart_stream:
                doc.add_picture(chart_stream, width=Inches(6))
        
        # Daily trends chart
        if MATPLOTLIB_AVAILABLE and daily:
            doc.add_paragraph()
            doc.add_heading("Daily Cost Trends", 2)
            chart_stream = self._create_trend_chart(daily, "Daily Cost Trends")
            if chart_stream:
                doc.add_picture(chart_stream, width=Inches(6))
        
        doc.add_paragraph()  # Spacing
    
    def _add_department_analysis_with_chart(self, doc: Document, data: Dict):
        """Add department-level analysis with chart."""
        doc.add_heading("Department Analysis", 1)
        
        sunburst = data.get('sunburst', {}).get('data', [])
        if not sunburst:
            doc.add_paragraph("No data available.")
            return
        
        # Group by department
        departments = {}
        for item in sunburst:
            if item.get('level') == 'drug':
                dept = item.get('parent_dept', 'Unknown')
                if dept not in departments:
                    departments[dept] = {'cost': 0, 'count': 0}
                departments[dept]['cost'] += item.get('value', 0)
                departments[dept]['count'] += 1
        
        rows = [["Department", "Total Cost", "Drugs Count", "Avg Cost per Drug"]]
        for dept in sorted(departments.keys(), key=lambda x: departments[x]['cost'], reverse=True):
            info = departments[dept]
            avg_cost = info['cost'] / info['count'] if info['count'] > 0 else 0
            rows.append([
                dept,
                f"${info['cost']:,.2f}",
                str(info['count']),
                f"${avg_cost:,.2f}"
            ])
        
        table = doc.add_table(rows=len(rows), cols=4)
        table.style = 'Light Grid Accent 1'
        
        # Format header
        hdr_cells = table.rows[0].cells
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].font.bold = True
        
        # Format data cells
        for row in table.rows[1:]:
            row.cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            row.cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            row.cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        doc.add_paragraph()  # Spacing
    
    def _add_category_analysis_with_chart(self, doc: Document, data: Dict):
        """Add drug category analysis with chart."""
        doc.add_heading("Drug Category Analysis", 1)
        
        sunburst = data.get('sunburst', {}).get('data', [])
        if not sunburst:
            doc.add_paragraph("No data available.")
            return
        
        # Group by category
        categories = {}
        for item in sunburst:
            if item.get('level') == 'category':
                category = item.get('name', 'Unknown')
                if category not in categories:
                    categories[category] = {'cost': 0, 'count': 0}
                categories[category]['cost'] += item.get('value', 0)
                # Count drugs under this category
                drug_count = sum(1 for drug in sunburst 
                               if drug.get('level') == 'drug' and drug.get('parent') == category)
                categories[category]['count'] = drug_count
        
        rows = [["Category", "Total Cost", "Drug Count", "% of Total"]]
        total_cost = sum(info['cost'] for info in categories.values())
        
        for category in sorted(categories.keys(), key=lambda x: categories[x]['cost'], reverse=True):
            info = categories[category]
            pct = (info['cost'] / total_cost * 100) if total_cost > 0 else 0
            rows.append([
                category,
                f"${info['cost']:,.2f}",
                str(info['count']),
                f"{pct:.1f}%"
            ])
        
        table = doc.add_table(rows=len(rows), cols=4)
        table.style = 'Light Grid Accent 1'
        
        # Format header
        hdr_cells = table.rows[0].cells
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].font.bold = True
        
        # Format data cells
        for row in table.rows[1:]:
            row.cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            row.cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            row.cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        # Add pie chart if matplotlib is available
        if MATPLOTLIB_AVAILABLE and categories:
            doc.add_paragraph()
            chart_stream = self._create_category_pie_chart(categories)
            if chart_stream:
                doc.add_picture(chart_stream, width=Inches(5))
        
        doc.add_paragraph()  # Spacing
    
    def _add_bubble_chart_analysis(self, doc: Document, data: Dict):
        """Add bubble chart analysis section."""
        doc.add_heading("Unit Price vs Quantity Analysis", 1)
        
        bubble_data = data.get('bubble_chart', {}).get('data', [])
        if not bubble_data:
            doc.add_paragraph("No bubble chart data available.")
            return
        
        # Summary
        p = doc.add_paragraph()
        p.add_run("This analysis shows the relationship between unit price and quantity for the top 200 drugs by frequency. ")
        p.add_run("The size of each bubble represents the frequency of prescription/usage.")
        
        doc.add_paragraph()
        
        # Statistics
        prices = [item.get('unit_price', 0) for item in bubble_data]
        quantities = [item.get('quantity', 0) for item in bubble_data]
        
        if prices and quantities:
            rows = [["Metric", "Unit Price", "Quantity"]]
            rows.append(["Average", f"${np.mean(prices):,.2f}", f"{np.mean(quantities):,.0f}"])
            rows.append(["Median", f"${np.median(prices):,.2f}", f"{np.median(quantities):,.0f}"])
            rows.append(["Max", f"${max(prices):,.2f}", f"{max(quantities):,.0f}"])
            rows.append(["Min", f"${min(prices):,.2f}", f"{min(quantities):,.0f}"])
            
            table = doc.add_table(rows=len(rows), cols=3)
            table.style = 'Light Grid Accent 1'
            
            # Format header
            hdr_cells = table.rows[0].cells
            for cell in hdr_cells:
                cell.paragraphs[0].runs[0].font.bold = True
            
            # Format data cells
            for row in table.rows[1:]:
                row.cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                row.cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        # Add bubble chart
        if MATPLOTLIB_AVAILABLE:
            doc.add_paragraph()
            chart_stream = self._create_bubble_chart(bubble_data)
            if chart_stream:
                doc.add_picture(chart_stream, width=Inches(6))
        
        doc.add_paragraph()  # Spacing
    
    def _add_detailed_sunburst_table(self, doc: Document, data: Dict):
        """Add detailed sunburst hierarchy table."""
        doc.add_heading("Detailed Cost Hierarchy", 1)
        
        sunburst = data.get('sunburst', {}).get('data', [])
        if not sunburst:
            doc.add_paragraph("No data available.")
            return
        
        # Limit to top items
        drug_items = [item for item in sunburst if item.get('level') == 'drug']
        drug_items = sorted(drug_items, key=lambda x: x.get('value', 0), reverse=True)[:50]
        
        rows = [["Department", "Category", "Drug", "Cost"]]
        for item in drug_items:
            rows.append([
                item.get('parent_dept', 'N/A'),
                item.get('parent', 'N/A'),
                item.get('name', 'N/A')[:40],
                f"${item.get('value', 0):,.2f}"
            ])
        
        table = doc.add_table(rows=len(rows), cols=4)
        table.style = 'Light Grid Accent 1'
        table.autofit = False
        
        # Set column widths
        widths = [Inches(1.5), Inches(1.5), Inches(2.5), Inches(1.2)]
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width
        
        # Format header
        hdr_cells = table.rows[0].cells
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].font.bold = True
        
        # Format data cells
        for row in table.rows[1:]:
            row.cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        doc.add_paragraph()  # Spacing
    
    def _add_footer(self, doc: Document):
        """Add footer to document."""
        doc.add_paragraph()
        
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = footer_para.add_run(
            "This report is generated automatically. "
            "For more information, please contact the analytics team."
        )
        run.font.size = Pt(9)
        run.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)
    
    # ========================================================================
    # Chart Generation Methods
    # ========================================================================
    
    def _create_department_chart(self, departments: Dict) -> BytesIO:
        """Create department cost bar chart."""
        if not MATPLOTLIB_AVAILABLE:
            return None
        
        try:
            # Sort and limit to top 10
            sorted_depts = sorted(departments.items(), key=lambda x: x[1][0], reverse=True)[:10]
            names = [dept[0] for dept in sorted_depts]
            costs = [dept[1][0] for dept in sorted_depts]
            
            fig, ax = plt.subplots(figsize=(10, 6))
            bars = ax.barh(names, costs, color=self.chart_color_palette[0])
            
            # Format
            ax.set_xlabel('Total Cost ($)', fontsize=12, fontweight='bold')
            ax.set_ylabel('Department', fontsize=12, fontweight='bold')
            ax.set_title('Top 10 Departments by Cost', fontsize=14, fontweight='bold')
            
            # Add value labels
            for i, (bar, cost) in enumerate(zip(bars, costs)):
                ax.text(cost, i, f' ${cost:,.0f}', va='center', fontsize=9)
            
            ax.grid(axis='x', alpha=0.3)
            plt.tight_layout()
            
            # Save to BytesIO
            stream = BytesIO()
            plt.savefig(stream, format='png', dpi=150, bbox_inches='tight')
            plt.close(fig)
            stream.seek(0)
            return stream
        except Exception as e:
            self.logger.error(f"Error creating department chart: {e}")
            return None
    
    def _create_top_drivers_chart(self, drivers: List[Dict]) -> BytesIO:
        """Create top cost drivers horizontal bar chart."""
        if not MATPLOTLIB_AVAILABLE:
            return None
        
        try:
            names = [d.get('drug_name', 'N/A')[:30] for d in drivers]
            costs = [d.get('total_cost', 0) for d in drivers]
            
            fig, ax = plt.subplots(figsize=(10, 8))
            bars = ax.barh(names, costs, color=self.chart_color_palette[1])
            
            # Format
            ax.set_xlabel('Total Cost ($)', fontsize=12, fontweight='bold')
            ax.set_ylabel('Drug Name', fontsize=12, fontweight='bold')
            ax.set_title('Top 10 Cost Drivers', fontsize=14, fontweight='bold')
            
            # Add value labels
            for i, (bar, cost) in enumerate(zip(bars, costs)):
                ax.text(cost, i, f' ${cost:,.0f}', va='center', fontsize=8)
            
            ax.grid(axis='x', alpha=0.3)
            plt.tight_layout()
            
            # Save to BytesIO
            stream = BytesIO()
            plt.savefig(stream, format='png', dpi=150, bbox_inches='tight')
            plt.close(fig)
            stream.seek(0)
            return stream
        except Exception as e:
            self.logger.error(f"Error creating top drivers chart: {e}")
            return None
    
    def _create_trend_chart(self, trends: List[Dict], title: str) -> BytesIO:
        """Create cost trend line chart."""
        if not MATPLOTLIB_AVAILABLE:
            return None
        
        try:
            dates = [t.get('date', '') for t in trends]
            costs = [t.get('total_cost', 0) for t in trends]
            
            fig, ax = plt.subplots(figsize=(10, 6))
            ax.plot(dates, costs, marker='o', linewidth=2, markersize=6, 
                   color=self.chart_color_palette[2])
            
            # Format
            ax.set_xlabel('Period', fontsize=12, fontweight='bold')
            ax.set_ylabel('Total Cost ($)', fontsize=12, fontweight='bold')
            ax.set_title(title, fontsize=14, fontweight='bold')
            
            # Rotate x labels
            plt.xticks(rotation=45, ha='right')
            
            # Format y-axis as currency
            ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f'${x:,.0f}'))
            
            ax.grid(alpha=0.3)
            plt.tight_layout()
            
            # Save to BytesIO
            stream = BytesIO()
            plt.savefig(stream, format='png', dpi=150, bbox_inches='tight')
            plt.close(fig)
            stream.seek(0)
            return stream
        except Exception as e:
            self.logger.error(f"Error creating trend chart: {e}")
            return None
    
    def _create_category_pie_chart(self, categories: Dict) -> BytesIO:
        """Create drug category pie chart."""
        if not MATPLOTLIB_AVAILABLE:
            return None
        
        try:
            # Sort and take top 8, group rest as "Other"
            sorted_cats = sorted(categories.items(), key=lambda x: x[1]['cost'], reverse=True)
            
            if len(sorted_cats) > 8:
                top_cats = sorted_cats[:8]
                other_cost = sum(cat[1]['cost'] for cat in sorted_cats[8:])
                labels = [cat[0] for cat in top_cats] + ['Other']
                costs = [cat[1]['cost'] for cat in top_cats] + [other_cost]
            else:
                labels = [cat[0] for cat in sorted_cats]
                costs = [cat[1]['cost'] for cat in sorted_cats]
            
            fig, ax = plt.subplots(figsize=(10, 8))
            colors = plt.cm.Set3(range(len(labels)))
            
            wedges, texts, autotexts = ax.pie(costs, labels=labels, autopct='%1.1f%%',
                                               startangle=90, colors=colors)
            
            # Format
            ax.set_title('Cost Distribution by Drug Category', fontsize=14, fontweight='bold')
            
            # Improve text readability
            for text in texts:
                text.set_fontsize(10)
            for autotext in autotexts:
                autotext.set_color('white')
                autotext.set_fontweight('bold')
                autotext.set_fontsize(9)
            
            plt.tight_layout()
            
            # Save to BytesIO
            stream = BytesIO()
            plt.savefig(stream, format='png', dpi=150, bbox_inches='tight')
            plt.close(fig)
            stream.seek(0)
            return stream
        except Exception as e:
            self.logger.error(f"Error creating category pie chart: {e}")
            return None
    
    def _create_bubble_chart(self, bubble_data: List[Dict]) -> BytesIO:
        """Create bubble chart for price vs quantity analysis."""
        if not MATPLOTLIB_AVAILABLE:
            return None
        
        try:
            # Limit to top 100 for readability
            bubble_data = bubble_data[:100]
            
            prices = [item.get('unit_price', 0) for item in bubble_data]
            quantities = [item.get('quantity', 0) for item in bubble_data]
            frequencies = [item.get('frequency', 1) for item in bubble_data]
            
            # Normalize bubble sizes
            max_freq = max(frequencies) if frequencies else 1
            sizes = [100 + (f / max_freq * 400) for f in frequencies]
            
            fig, ax = plt.subplots(figsize=(12, 8))
            scatter = ax.scatter(prices, quantities, s=sizes, alpha=0.5, 
                               c=range(len(prices)), cmap='viridis', edgecolors='w', linewidth=0.5)
            
            # Format
            ax.set_xlabel('Unit Price ($)', fontsize=12, fontweight='bold')
            ax.set_ylabel('Quantity', fontsize=12, fontweight='bold')
            ax.set_title('Unit Price vs Quantity (Bubble Size = Frequency)', 
                        fontsize=14, fontweight='bold')
            
            # Format axes
            ax.xaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f'${x:,.0f}'))
            ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f'{x:,.0f}'))
            
            ax.grid(alpha=0.3)
            plt.tight_layout()
            
            # Save to BytesIO
            stream = BytesIO()
            plt.savefig(stream, format='png', dpi=150, bbox_inches='tight')
            plt.close(fig)
            stream.seek(0)
            return stream
        except Exception as e:
            self.logger.error(f"Error creating bubble chart: {e}")
            return None
    
    # ========================================================================
    # PDF Helper Methods
    # ========================================================================
    
    def _generate_summary_text(self, data: Dict, filters: Dict) -> str:
        """Generate summary text for the report."""
        sunburst = data.get('sunburst', {}).get('data', [])
        drivers = data.get('top_cost_drivers', {}).get('data', [])
        
        # Calculate total cost and unique drugs from drug-level items
        drug_items = [item for item in sunburst if item.get('level') == 'drug']
        total_cost = sum(item.get('value', 0) for item in drug_items)
        unique_drugs = len(drug_items)
        
        # If sunburst is empty or has no cost, try to use top drivers data
        if (total_cost == 0 or unique_drugs == 0) and drivers:
            total_cost = sum(d.get('total_cost', 0) for d in drivers)
            unique_drugs = len(drivers)
        
        top_drug = drivers[0].get('drug_name', 'N/A') if drivers else 'N/A'
        top_drug_cost = drivers[0].get('total_cost', 0) if drivers else 0
        
        start_date = filters.get('start_date', 'N/A')
        end_date = filters.get('end_date', 'N/A')
        
        summary = f"""
        This report analyzes pharmaceutical costs from {start_date} to {end_date}. 
        The total cost during this period was <b>${total_cost:,.2f}</b> across <b>{unique_drugs}</b> unique drugs. 
        The top cost driver was <b>{top_drug}</b> with a cost of <b>${top_drug_cost:,.2f}</b>.
        This report provides detailed insights into cost distribution across departments, drug categories, 
        trends over time, and price-quantity relationships to support informed decision-making.
        """
        return summary.strip()
    
    def _generate_bubble_summary(self, data: Dict) -> str:
        """Generate summary text for bubble chart analysis."""
        bubble_data = data.get('bubble_chart', {}).get('data', [])
        
        if not bubble_data:
            return "No price-quantity data available for analysis."
        
        prices = [item.get('unit_price', 0) for item in bubble_data]
        quantities = [item.get('quantity', 0) for item in bubble_data]
        
        avg_price = np.mean(prices) if prices else 0
        avg_qty = np.mean(quantities) if quantities else 0
        
        summary = f"""
        This analysis examines the relationship between unit price and quantity for the top {len(bubble_data)} drugs by frequency. 
        The average unit price is <b>${avg_price:,.2f}</b> with an average quantity of <b>{avg_qty:,.0f}</b> units. 
        Bubble size represents prescription/usage frequency, helping identify high-volume, high-cost items 
        that may be targets for cost optimization initiatives.
        """
        return summary.strip()
    
    def _create_filters_table(self, filters: Dict) -> Table:
        """Create filters table for PDF."""
        departments = filters.get('departments') or []
        if isinstance(departments, (list, tuple)):
            departments_str = ', '.join(map(str, departments)) if departments else "All"
        else:
            departments_str = "All"
            
        drug_categories = filters.get('drug_categories') or []
        if isinstance(drug_categories, (list, tuple)):
            categories_str = ', '.join(map(str, drug_categories)) if drug_categories else "All"
        else:
            categories_str = "All"
            
        price_min = filters.get('price_min') or 0
        price_max = filters.get('price_max') or 'Unlimited'
        
        data = [
            ["Analysis Period", f"{filters.get('start_date', 'N/A')} to {filters.get('end_date', 'N/A')}"],
            ["Departments", departments_str],
            ["Price Range", f"${price_min} - ${price_max}"],
            ["Drug Categories", categories_str],
        ]
        
        table = Table(data, colWidths=[1.5*inch, 4*inch])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#1f4788')),
            ('TEXTCOLOR', (0, 0), (0, -1), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('TOPPADDING', (0, 0), (-1, -1), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey),
        ]))
        
        return table
    
    def _create_cost_overview_table(self, data: Dict) -> Table:
        """Create cost overview table for PDF."""
        sunburst = data.get('sunburst', {}).get('data', [])
        
        if not sunburst:
            return None
        
        table_data = [["Department", "Total Cost", "Categories", "Drugs"]]
        
        # Group by department - try multiple approaches
        departments = {}
        
        # First, try to get departments directly
        for item in sunburst:
            if item.get('level') == 'department':
                dept_name = item.get('name', 'Unknown')
                dept_cost = item.get('value', 0)
                cat_count = sum(1 for sub in sunburst if sub.get('parent') == dept_name and sub.get('level') == 'category')
                drug_count = sum(1 for sub in sunburst if sub.get('parent_dept') == dept_name and sub.get('level') == 'drug')
                departments[dept_name] = {'cost': dept_cost, 'categories': cat_count, 'drugs': drug_count}
        
        # If no departments found, aggregate from drugs
        if not departments:
            for item in sunburst:
                if item.get('level') == 'drug':
                    dept = item.get('parent_dept', 'Unknown')
                    if dept not in departments:
                        departments[dept] = {'cost': 0, 'categories': set(), 'drugs': 0}
                    departments[dept]['cost'] += item.get('value', 0)
                    departments[dept]['categories'].add(item.get('parent', 'Unknown'))
                    departments[dept]['drugs'] += 1
            
            # Convert sets to counts
            for dept in departments:
                departments[dept]['categories'] = len(departments[dept]['categories'])
        
        # If still no data, return None
        if not departments:
            return None
        
        for dept in sorted(departments.keys(), key=lambda x: departments[x]['cost'], reverse=True)[:10]:
            info = departments[dept]
            cat_count = info['categories'] if isinstance(info['categories'], int) else len(info['categories'])
            table_data.append([
                dept,
                f"${info['cost']:,.2f}",
                str(cat_count),
                str(info['drugs'])
            ])
        
        # If only header, return None
        if len(table_data) == 1:
            return None
        
        table = Table(table_data, colWidths=[2*inch, 1.3*inch, 1*inch, 0.8*inch])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f4788')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('ALIGN', (1, 1), (-1, -1), 'RIGHT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f0f0f0')]),
        ]))
        
        return table
    
    def _create_top_drivers_table(self, data: Dict) -> Table:
        """Create top cost drivers table for PDF."""
        drivers = data.get('top_cost_drivers', {}).get('data', [])
        
        if not drivers:
            return None
        
        table_data = [["Rank", "Drug Name", "Total Cost", "Units"]]
        
        for idx, driver in enumerate(drivers[:20], 1):
            table_data.append([
                str(idx),
                driver.get('drug_name', 'N/A')[:35],
                f"${driver.get('total_cost', 0):,.2f}",
                str(driver.get('units', 0))
            ])
        
        table = Table(table_data, colWidths=[0.4*inch, 2.8*inch, 1.2*inch, 0.8*inch])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f4788')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('ALIGN', (0, 0), (0, -1), 'CENTER'),
            ('ALIGN', (2, 1), (-1, -1), 'RIGHT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f0f0f0')]),
        ]))
        
        return table
    
    def _create_monthly_trends_table(self, data: Dict) -> Table:
        """Create monthly trends table for PDF."""
        monthly = data.get('cost_trends', {}).get('monthly', {}).get('data', [])
        
        if not monthly:
            return None
        
        table_data = [["Period", "Total Cost", "Change %"]]
        prev_cost = None
        
        for trend in monthly[:12]:
            period = trend.get('date', 'N/A')
            cost = trend.get('total_cost', 0)
            
            change = "—"
            if prev_cost is not None and prev_cost > 0:
                pct_change = ((cost - prev_cost) / prev_cost) * 100
                change = f"{pct_change:+.1f}%"
            
            table_data.append([period, f"${cost:,.2f}", change])
            prev_cost = cost
        
        table = Table(table_data, colWidths=[1.5*inch, 1.5*inch, 1*inch])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f4788')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('ALIGN', (1, 1), (-1, -1), 'RIGHT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f0f0f0')]),
        ]))
        
        return table
    
    def _create_department_table(self, data: Dict) -> Table:
        """Create department analysis table for PDF."""
        sunburst = data.get('sunburst', {}).get('data', [])
        
        if not sunburst:
            return None
        
        # Group by department from drug-level data
        departments = {}
        for item in sunburst:
            if item.get('level') == 'drug':
                dept = item.get('parent_dept', 'Unknown')
                if dept != 'Unknown':  # Skip unknown departments
                    if dept not in departments:
                        departments[dept] = {'cost': 0, 'count': 0}
                    departments[dept]['cost'] += item.get('value', 0)
                    departments[dept]['count'] += 1
        
        # If no departments found, return None
        if not departments:
            return None
        
        table_data = [["Department", "Total Cost", "Drugs", "Avg Cost"]]
        
        for dept in sorted(departments.keys(), key=lambda x: departments[x]['cost'], reverse=True)[:15]:
            info = departments[dept]
            avg_cost = info['cost'] / info['count'] if info['count'] > 0 else 0
            table_data.append([
                dept,
                f"${info['cost']:,.2f}",
                str(info['count']),
                f"${avg_cost:,.2f}"
            ])
        
        # If only header, return None
        if len(table_data) == 1:
            return None
        
        table = Table(table_data, colWidths=[2*inch, 1.3*inch, 0.7*inch, 1*inch])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f4788')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('ALIGN', (1, 1), (-1, -1), 'RIGHT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f0f0f0')]),
        ]))
        
        return table
    
    def _create_category_table(self, data: Dict) -> Table:
        """Create category analysis table for PDF."""
        sunburst = data.get('sunburst', {}).get('data', [])
        
        if not sunburst:
            return None
        
        # Group by category - try multiple approaches
        categories = {}
        
        # First approach: Get categories directly from level='category'
        for item in sunburst:
            if item.get('level') == 'category':
                category = item.get('name', 'Unknown')
                if category != 'Unknown':
                    if category not in categories:
                        categories[category] = {'cost': 0, 'count': 0}
                    categories[category]['cost'] += item.get('value', 0)
                    # Count drugs under this category
                    drug_count = sum(1 for drug in sunburst 
                                   if drug.get('level') == 'drug' and drug.get('parent') == category)
                    categories[category]['count'] = drug_count
        
        # Second approach: If no categories found, aggregate from drug-level data
        if not categories:
            for item in sunburst:
                if item.get('level') == 'drug':
                    category = item.get('parent_cat') or item.get('parent', 'Unknown')
                    if category != 'Unknown':
                        if category not in categories:
                            categories[category] = {'cost': 0, 'count': 0}
                        categories[category]['cost'] += item.get('value', 0)
                        categories[category]['count'] += 1
        
        # If still no categories, return None
        if not categories:
            return None
        
        table_data = [["Category", "Total Cost", "Drugs", "% of Total"]]
        total_cost = sum(info['cost'] for info in categories.values())
        
        if total_cost == 0:
            return None
        
        for category in sorted(categories.keys(), key=lambda x: categories[x]['cost'], reverse=True)[:15]:
            info = categories[category]
            pct = (info['cost'] / total_cost * 100) if total_cost > 0 else 0
            table_data.append([
                category,
                f"${info['cost']:,.2f}",
                str(info['count']),
                f"{pct:.1f}%"
            ])
        
        # If only header, return None
        if len(table_data) == 1:
            return None
        
        table = Table(table_data, colWidths=[2*inch, 1.3*inch, 0.7*inch, 0.8*inch])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f4788')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('ALIGN', (1, 1), (-1, -1), 'RIGHT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f0f0f0')]),
        ]))
        
        return table
    
    # ========================================================================
    # PDF Chart Generation Methods
    # ========================================================================
    
    def _create_cost_overview_chart_pdf(self, data: Dict) -> Image:
        """Create cost overview chart for PDF."""
        if not MATPLOTLIB_AVAILABLE:
            return None
        
        try:
            sunburst = data.get('sunburst', {}).get('data', [])
            
            # Group by department
            departments = {}
            for item in sunburst:
                if item.get('level') == 'department':
                    dept_name = item.get('name', 'Unknown')
                    dept_cost = item.get('value', 0)
                    departments[dept_name] = dept_cost
            
            if not departments:
                return None
            
            chart_stream = self._create_department_chart(
                {k: (v, 0) for k, v in departments.items()}
            )
            
            if chart_stream:
                return Image(chart_stream, width=5*inch, height=3*inch)
            return None
        except Exception as e:
            self.logger.error(f"Error creating overview chart for PDF: {e}")
            return None
    
    def _create_top_drivers_chart_pdf(self, data: Dict) -> Image:
        """Create top drivers chart for PDF."""
        if not MATPLOTLIB_AVAILABLE:
            return None
        
        try:
            drivers = data.get('top_cost_drivers', {}).get('data', [])[:10]
            if not drivers:
                return None
            
            chart_stream = self._create_top_drivers_chart(drivers)
            
            if chart_stream:
                return Image(chart_stream, width=5*inch, height=4*inch)
            return None
        except Exception as e:
            self.logger.error(f"Error creating drivers chart for PDF: {e}")
            return None
    
    def _create_cost_trends_chart_pdf(self, data: Dict) -> Image:
        """Create cost trends chart for PDF."""
        if not MATPLOTLIB_AVAILABLE:
            return None
        
        try:
            monthly = data.get('cost_trends', {}).get('monthly', {}).get('data', [])
            if not monthly:
                return None
            
            chart_stream = self._create_trend_chart(monthly, "Monthly Cost Trends")
            
            if chart_stream:
                return Image(chart_stream, width=5.5*inch, height=3.5*inch)
            return None
        except Exception as e:
            self.logger.error(f"Error creating trends chart for PDF: {e}")
            return None
    
    def _create_bubble_chart_pdf(self, data: Dict) -> Image:
        """Create bubble chart for PDF."""
        if not MATPLOTLIB_AVAILABLE:
            return None
        
        try:
            bubble_data = data.get('bubble_chart', {}).get('data', [])
            if not bubble_data:
                return None
            
            chart_stream = self._create_bubble_chart(bubble_data)
            
            if chart_stream:
                return Image(chart_stream, width=6*inch, height=4*inch)
            return None
        except Exception as e:
            self.logger.error(f"Error creating bubble chart for PDF: {e}")
            return None
    
    # ========================================================================
    # HOSPITAL STAY REPORT METHODS
    # ========================================================================
    
    def generate_hospital_stay_docx_report(self, hospital_stay_data: Dict, filters: Dict) -> BinaryIO:
        """
        Generate a comprehensive DOCX report with hospital stay analysis data and charts.
        
        Args:
            hospital_stay_data: Data from hospital stay analysis API
            filters: Applied filters for the analysis
            
        Returns:
            BytesIO object containing the DOCX file
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for DOCX report generation. "
                "Please install it: pip install python-docx"
            )
        
        doc = Document()
        
        # Add styles
        self._setup_styles(doc)
        
        # Title
        self._add_title(doc, "Hospital Stay Duration Analysis Report")
        
        # Executive Summary
        self._add_stay_executive_summary(doc, hospital_stay_data, filters)
        
        # Filters Applied
        self._add_stay_filters_section(doc, filters)
        
        # Statistics Overview
        self._add_stay_statistics(doc, hospital_stay_data)
        
        # Distribution Analysis with Chart
        self._add_stay_distribution(doc, hospital_stay_data)
        
        # Department Analysis with Chart
        self._add_stay_by_department(doc, hospital_stay_data)
        
        # Monthly Trends with Chart
        self._add_stay_trends(doc, hospital_stay_data)
        
        # Top Patient Stays Table
        self._add_patient_stays_table(doc, hospital_stay_data)
        
        # Footer
        self._add_footer(doc)
        
        # Save to BytesIO
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output
    
    def generate_hospital_stay_pdf_report(self, hospital_stay_data: Dict, filters: Dict) -> BinaryIO:
        """
        Generate a comprehensive PDF report with hospital stay analysis data and charts.
        
        Args:
            hospital_stay_data: Data from hospital stay analysis API
            filters: Applied filters for the analysis
            
        Returns:
            BytesIO object containing the PDF file
        """
        if not REPORTLAB_AVAILABLE:
            raise ImportError(
                "reportlab is required for PDF report generation. "
                "Please install it: pip install reportlab"
            )
        
        output = BytesIO()
        doc = SimpleDocTemplate(output, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
        
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1f4788'),
            spaceAfter=30,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#1f4788'),
            spaceAfter=12,
            spaceBefore=12,
            fontName='Helvetica-Bold'
        )
        
        subheading_style = ParagraphStyle(
            'CustomSubHeading',
            parent=styles['Heading3'],
            fontSize=12,
            textColor=colors.HexColor('#2874A6'),
            spaceAfter=8,
            spaceBefore=8,
            fontName='Helvetica-Bold'
        )
        
        # Title
        story.append(Paragraph("Hospital Stay Duration Analysis Report", title_style))
        story.append(Spacer(1, 0.2*inch))
        
        # Date and filters info
        report_date = datetime.now().strftime("%B %d, %Y")
        story.append(Paragraph(f"<b>Report Date:</b> {report_date}", styles['Normal']))
        story.append(Spacer(1, 0.1*inch))
        
        # Executive Summary
        story.append(Paragraph("Executive Summary", heading_style))
        summary_text = self._generate_stay_summary_text(hospital_stay_data, filters)
        story.append(Paragraph(summary_text, styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        
        # Filters Applied
        story.append(Paragraph("Analysis Filters", heading_style))
        filters_table = self._create_stay_filters_table(filters)
        story.append(filters_table)
        story.append(Spacer(1, 0.2*inch))
        
        # Statistics Table
        story.append(Paragraph("Stay Duration Statistics", heading_style))
        stats_table = self._create_stay_statistics_table(hospital_stay_data)
        story.append(stats_table)
        story.append(Spacer(1, 0.2*inch))
        
        story.append(PageBreak())
        
        # Distribution Analysis
        story.append(Paragraph("Stay Duration Distribution", heading_style))
        dist_table = self._create_stay_distribution_table(hospital_stay_data)
        if dist_table:
            story.append(dist_table)
        story.append(Spacer(1, 0.2*inch))
        
        # Distribution Chart
        if MATPLOTLIB_AVAILABLE:
            dist_chart = self._create_stay_distribution_chart_pdf(hospital_stay_data)
            if dist_chart:
                story.append(dist_chart)
                story.append(Spacer(1, 0.2*inch))
        
        story.append(PageBreak())
        
        # Department Analysis
        story.append(Paragraph("Stay Duration by Department", heading_style))
        dept_table = self._create_stay_department_table(hospital_stay_data)
        if dept_table:
            story.append(dept_table)
        story.append(Spacer(1, 0.2*inch))
        
        # Department Chart
        if MATPLOTLIB_AVAILABLE:
            dept_chart = self._create_stay_department_chart_pdf(hospital_stay_data)
            if dept_chart:
                story.append(dept_chart)
                story.append(Spacer(1, 0.2*inch))
        
        story.append(PageBreak())
        
        # Monthly Trends
        story.append(Paragraph("Monthly Trends", heading_style))
        trends_table = self._create_stay_trends_table(hospital_stay_data)
        if trends_table:
            story.append(trends_table)
        story.append(Spacer(1, 0.2*inch))
        
        # Trends Chart
        if MATPLOTLIB_AVAILABLE:
            trends_chart = self._create_stay_trends_chart_pdf(hospital_stay_data)
            if trends_chart:
                story.append(trends_chart)
                story.append(Spacer(1, 0.2*inch))
        
        story.append(PageBreak())
        
        # Top Patient Stays
        story.append(Paragraph("Longest Patient Stays", heading_style))
        patients_table = self._create_patient_stays_table_pdf(hospital_stay_data)
        if patients_table:
            story.append(patients_table)
        
        # Build PDF
        doc.build(story)
        output.seek(0)
        return output
    
    # ========================================================================
    # Hospital Stay DOCX Helper Methods
    # ========================================================================
    
    def _add_stay_executive_summary(self, doc: Document, data: Dict, filters: Dict):
        """Add executive summary for hospital stay report."""
        doc.add_heading("Executive Summary", 1)
        
        # Get statistics
        statistics = data.get('statistics', {}).get('data', {})
        
        # Summary text
        p = doc.add_paragraph()
        p.add_run(f"This comprehensive report analyzes hospital stay durations from ").font.size = Pt(11)
        run = p.add_run(f"{filters.get('start_date', 'N/A')} to {filters.get('end_date', 'N/A')}")
        run.bold = True
        run.font.size = Pt(11)
        p.add_run(".").font.size = Pt(11)
        
        # Key metrics
        doc.add_paragraph()
        doc.add_heading("Key Statistics", 2)
        
        metrics_data = [
            ["Metric", "Value"],
            ["Total Patients", f"{statistics.get('total_patients', 0):,}"],
            ["Average Stay", f"{statistics.get('average_stay_days', 0):.1f} days"],
            ["Median Stay", f"{statistics.get('median_stay_days', 0):.1f} days"],
            ["Shortest Stay", f"{statistics.get('min_stay_days', 0):.1f} days"],
            ["Longest Stay", f"{statistics.get('max_stay_days', 0):.1f} days"],
            ["Standard Deviation", f"{statistics.get('std_dev_days', 0):.1f} days"],
        ]
        
        table = doc.add_table(rows=len(metrics_data), cols=2)
        table.style = 'Light Grid Accent 1'
        
        for i, row_data in enumerate(metrics_data):
            row = table.rows[i]
            row.cells[0].text = row_data[0]
            row.cells[1].text = row_data[1]
            if i == 0:  # Header
                for cell in row.cells:
                    cell.paragraphs[0].runs[0].font.bold = True
            else:
                row.cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        doc.add_paragraph()  # Spacing
    
    def _add_stay_filters_section(self, doc: Document, filters: Dict):
        """Add applied filters section for hospital stay report."""
        doc.add_heading("Analysis Parameters", 1)
        
        # Handle None values safely
        departments = filters.get('departments') or []
        departments_str = ', '.join(map(str, departments)) if (isinstance(departments, (list, tuple)) and departments) else "All"
        
        min_stay = filters.get('min_stay_days') or 'No minimum'
        max_stay = filters.get('max_stay_days') or 'No maximum'
        
        rows = [["Parameter", "Value"]]
        rows.append(["Start Date", filters.get('start_date', 'N/A')])
        rows.append(["End Date", filters.get('end_date', 'N/A')])
        rows.append(["Departments", departments_str])
        rows.append(["Stay Duration Range", f"{min_stay} - {max_stay} days"])
        
        table = doc.add_table(rows=len(rows), cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Format header
        hdr_cells = table.rows[0].cells
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()  # Spacing
    
    def _add_stay_statistics(self, doc: Document, data: Dict):
        """Add statistics overview section."""
        doc.add_heading("Statistical Summary", 1)
        
        statistics = data.get('statistics', {}).get('data', {})
        
        if not statistics:
            doc.add_paragraph("No statistics available.")
            return
        
        # Create detailed statistics table
        rows = [
            ["Statistic", "Value", "Description"],
            ["Total Patients", f"{statistics.get('total_patients', 0):,}", "Number of patients analyzed"],
            ["Average Stay", f"{statistics.get('average_stay_days', 0):.2f} days", "Mean stay duration"],
            ["Median Stay", f"{statistics.get('median_stay_days', 0):.2f} days", "Middle value of stay durations"],
            ["Minimum Stay", f"{statistics.get('min_stay_days', 0):.2f} days", "Shortest recorded stay"],
            ["Maximum Stay", f"{statistics.get('max_stay_days', 0):.2f} days", "Longest recorded stay"],
            ["Standard Deviation", f"{statistics.get('std_dev_days', 0):.2f} days", "Measure of variability"],
        ]
        
        table = doc.add_table(rows=len(rows), cols=3)
        table.style = 'Light Grid Accent 1'
        table.autofit = False
        
        # Set column widths
        widths = [Inches(2.0), Inches(1.5), Inches(2.5)]
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width
        
        # Format header
        hdr_cells = table.rows[0].cells
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].font.bold = True
        
        # Format data cells
        for row in table.rows[1:]:
            row.cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        doc.add_paragraph()  # Spacing
    
    def _add_stay_distribution(self, doc: Document, data: Dict):
        """Add stay duration distribution section with chart."""
        doc.add_heading("Stay Duration Distribution", 1)
        
        distribution = data.get('distribution', {}).get('data', [])
        
        if not distribution:
            doc.add_paragraph("No distribution data available.")
            return
        
        # Distribution table
        rows = [["Duration Range", "Patient Count", "Percentage"]]
        total_patients = sum(item.get('patient_count', 0) for item in distribution)
        
        for item in distribution:
            count = item.get('patient_count', 0)
            pct = (count / total_patients * 100) if total_patients > 0 else 0
            rows.append([
                item.get('stay_duration_range', 'N/A'),
                str(count),
                f"{pct:.1f}%"
            ])
        
        table = doc.add_table(rows=len(rows), cols=3)
        table.style = 'Light Grid Accent 1'
        
        # Format header
        hdr_cells = table.rows[0].cells
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].font.bold = True
        
        # Format data cells
        for row in table.rows[1:]:
            row.cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            row.cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        # Add chart if matplotlib is available
        if MATPLOTLIB_AVAILABLE and distribution:
            doc.add_paragraph()
            chart_stream = self._create_stay_distribution_chart(distribution)
            if chart_stream:
                doc.add_picture(chart_stream, width=Inches(6))
        
        doc.add_paragraph()  # Spacing
    
    def _add_stay_by_department(self, doc: Document, data: Dict):
        """Add stay duration by department section with chart."""
        doc.add_heading("Stay Duration by Department", 1)
        
        by_department = data.get('by_department', {}).get('data', [])
        
        if not by_department:
            doc.add_paragraph("No department data available.")
            return
        
        # Department table
        rows = [["Department", "Avg Stay (days)", "Patient Count"]]
        
        for dept in sorted(by_department, key=lambda x: x.get('average_stay_days', 0), reverse=True):
            rows.append([
                dept.get('department_name', f"Dept {dept.get('department_id', 'N/A')}"),
                f"{dept.get('average_stay_days', 0):.1f}",
                str(dept.get('patient_count', 0))
            ])
        
        table = doc.add_table(rows=len(rows), cols=3)
        table.style = 'Light Grid Accent 1'
        
        # Format header
        hdr_cells = table.rows[0].cells
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].font.bold = True
        
        # Format data cells
        for row in table.rows[1:]:
            row.cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            row.cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        # Add chart if matplotlib is available
        if MATPLOTLIB_AVAILABLE and by_department:
            doc.add_paragraph()
            chart_stream = self._create_stay_department_chart(by_department)
            if chart_stream:
                doc.add_picture(chart_stream, width=Inches(6))
        
        doc.add_paragraph()  # Spacing
    
    def _add_stay_trends(self, doc: Document, data: Dict):
        """Add monthly trends section with chart."""
        doc.add_heading("Monthly Stay Duration Trends", 1)
        
        trends = data.get('trends', {}).get('monthly', {}).get('data', [])
        
        if not trends:
            doc.add_paragraph("No trend data available.")
            return
        
        # Trends table
        rows = [["Period", "Avg Stay (days)", "Patient Count", "Change %"]]
        prev_avg = None
        
        for trend in trends[:12]:  # Last 12 months
            period = trend.get('period', 'N/A')
            avg_stay = trend.get('average_stay_days', 0)
            count = trend.get('patient_count', 0)
            
            change = "—"
            if prev_avg is not None and prev_avg > 0:
                pct_change = ((avg_stay - prev_avg) / prev_avg) * 100
                change = f"{pct_change:+.1f}%"
            
            rows.append([period, f"{avg_stay:.1f}", str(count), change])
            prev_avg = avg_stay
        
        table = doc.add_table(rows=len(rows), cols=4)
        table.style = 'Light Grid Accent 1'
        
        # Format header
        hdr_cells = table.rows[0].cells
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].font.bold = True
        
        # Format data cells
        for row in table.rows[1:]:
            row.cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            row.cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            row.cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        # Add chart if matplotlib is available
        if MATPLOTLIB_AVAILABLE and trends:
            doc.add_paragraph()
            chart_stream = self._create_stay_trends_chart(trends)
            if chart_stream:
                doc.add_picture(chart_stream, width=Inches(6))
        
        doc.add_paragraph()  # Spacing
    
    def _add_patient_stays_table(self, doc: Document, data: Dict):
        """Add top patient stays table."""
        doc.add_heading("Longest Patient Stays (Top 50)", 1)
        
        patient_stays = data.get('patient_stays', {}).get('data', [])
        total_count = data.get('patient_stays', {}).get('config', {}).get('total_count', 0)
        
        if not patient_stays:
            doc.add_paragraph("No patient stay data available.")
            return
        
        # Info paragraph
        doc.add_paragraph(f"Displaying top 50 of {total_count:,} total patient stays, sorted by duration.")
        doc.add_paragraph()
        
        # Patient stays table (top 50)
        rows = [["Rank", "Patient ID", "Admission Date", "Discharge Date", "Stay (days)", "Department"]]
        
        for idx, stay in enumerate(patient_stays[:50], 1):
            rows.append([
                str(idx),
                str(stay.get('patient_id', 'N/A')),
                stay.get('admission_date', 'N/A'),
                stay.get('discharge_date', 'N/A'),
                f"{stay.get('stay_days', 0):.1f}",
                stay.get('department_name', f"Dept {stay.get('department_id', 'N/A')}")
            ])
        
        table = doc.add_table(rows=len(rows), cols=6)
        table.style = 'Light Grid Accent 1'
        table.autofit = False
        
        # Set column widths
        widths = [Inches(0.5), Inches(1.0), Inches(1.2), Inches(1.2), Inches(0.9), Inches(1.2)]
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width
        
        # Format header
        hdr_cells = table.rows[0].cells
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(9)
        
        # Format data cells
        for row in table.rows[1:]:
            for cell in row.cells:
                cell.paragraphs[0].runs[0].font.size = Pt(8)
            row.cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row.cells[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        doc.add_paragraph()  # Spacing
    
    # ========================================================================
    # Hospital Stay Chart Generation Methods
    # ========================================================================
    
    def _create_stay_distribution_chart(self, distribution: List[Dict]) -> BytesIO:
        """Create stay duration distribution histogram."""
        if not MATPLOTLIB_AVAILABLE:
            return None
        
        try:
            ranges = [item.get('stay_duration_range', '') for item in distribution]
            counts = [item.get('patient_count', 0) for item in distribution]
            
            fig, ax = plt.subplots(figsize=(10, 6))
            bars = ax.bar(range(len(ranges)), counts, color=self.chart_color_palette[3])
            
            # Format
            ax.set_xlabel('Stay Duration Range', fontsize=12, fontweight='bold')
            ax.set_ylabel('Patient Count', fontsize=12, fontweight='bold')
            ax.set_title('Distribution of Hospital Stay Durations', fontsize=14, fontweight='bold')
            ax.set_xticks(range(len(ranges)))
            ax.set_xticklabels(ranges, rotation=45, ha='right')
            
            # Add value labels on bars
            for bar in bars:
                height = bar.get_height()
                ax.text(bar.get_x() + bar.get_width()/2., height,
                       f'{int(height)}',
                       ha='center', va='bottom', fontsize=9)
            
            ax.grid(axis='y', alpha=0.3)
            plt.tight_layout()
            
            # Save to BytesIO
            stream = BytesIO()
            plt.savefig(stream, format='png', dpi=150, bbox_inches='tight')
            plt.close(fig)
            stream.seek(0)
            return stream
        except Exception as e:
            self.logger.error(f"Error creating distribution chart: {e}")
            return None
    
    def _create_stay_department_chart(self, by_department: List[Dict]) -> BytesIO:
        """Create department comparison bar chart."""
        if not MATPLOTLIB_AVAILABLE:
            return None
        
        try:
            # Sort by average stay
            sorted_depts = sorted(by_department, key=lambda x: x.get('average_stay_days', 0), reverse=True)[:10]
            
            names = [d.get('department_name', f"Dept {d.get('department_id', 'N/A')}") for d in sorted_depts]
            avg_stays = [d.get('average_stay_days', 0) for d in sorted_depts]
            
            fig, ax = plt.subplots(figsize=(10, 6))
            bars = ax.barh(names, avg_stays, color=self.chart_color_palette[0])
            
            # Format
            ax.set_xlabel('Average Stay Duration (days)', fontsize=12, fontweight='bold')
            ax.set_ylabel('Department', fontsize=12, fontweight='bold')
            ax.set_title('Top 10 Departments by Average Stay Duration', fontsize=14, fontweight='bold')
            
            # Add value labels
            for i, (bar, stay) in enumerate(zip(bars, avg_stays)):
                ax.text(stay, i, f' {stay:.1f} days', va='center', fontsize=9)
            
            ax.grid(axis='x', alpha=0.3)
            plt.tight_layout()
            
            # Save to BytesIO
            stream = BytesIO()
            plt.savefig(stream, format='png', dpi=150, bbox_inches='tight')
            plt.close(fig)
            stream.seek(0)
            return stream
        except Exception as e:
            self.logger.error(f"Error creating department chart: {e}")
            return None
    
    def _create_stay_trends_chart(self, trends: List[Dict]) -> BytesIO:
        """Create monthly trends line chart."""
        if not MATPLOTLIB_AVAILABLE:
            return None
        
        try:
            periods = [t.get('period', '') for t in trends]
            avg_stays = [t.get('average_stay_days', 0) for t in trends]
            
            fig, ax = plt.subplots(figsize=(10, 6))
            ax.plot(periods, avg_stays, marker='o', linewidth=2, markersize=8, 
                   color=self.chart_color_palette[2])
            
            # Format
            ax.set_xlabel('Period', fontsize=12, fontweight='bold')
            ax.set_ylabel('Average Stay Duration (days)', fontsize=12, fontweight='bold')
            ax.set_title('Monthly Trends in Hospital Stay Duration', fontsize=14, fontweight='bold')
            
            # Rotate x labels
            plt.xticks(rotation=45, ha='right')
            
            ax.grid(alpha=0.3)
            plt.tight_layout()
            
            # Save to BytesIO
            stream = BytesIO()
            plt.savefig(stream, format='png', dpi=150, bbox_inches='tight')
            plt.close(fig)
            stream.seek(0)
            return stream
        except Exception as e:
            self.logger.error(f"Error creating trends chart: {e}")
            return None
    
    # ========================================================================
    # Hospital Stay PDF Helper Methods
    # ========================================================================
    
    def _generate_stay_summary_text(self, data: Dict, filters: Dict) -> str:
        """Generate summary text for hospital stay report."""
        statistics = data.get('statistics', {}).get('data', {})
        
        total_patients = statistics.get('total_patients', 0)
        avg_stay = statistics.get('average_stay_days', 0)
        max_stay = statistics.get('max_stay_days', 0)
        
        start_date = filters.get('start_date', 'N/A')
        end_date = filters.get('end_date', 'N/A')
        
        summary = f"""
        This report analyzes hospital stay durations from {start_date} to {end_date}. 
        A total of <b>{total_patients:,} patients</b> were analyzed during this period. 
        The average stay duration was <b>{avg_stay:.1f} days</b>, with the longest stay being <b>{max_stay:.1f} days</b>.
        This report provides insights into stay duration patterns, department variations, and trends over time
        to support operational planning and resource allocation.
        """
        return summary.strip()
    
    def _create_stay_filters_table(self, filters: Dict) -> Table:
        """Create filters table for hospital stay PDF."""
        departments = filters.get('departments') or []
        if isinstance(departments, (list, tuple)):
            departments_str = ', '.join(map(str, departments)) if departments else "All"
        else:
            departments_str = "All"
            
        min_stay = filters.get('min_stay_days') or 'No minimum'
        max_stay = filters.get('max_stay_days') or 'No maximum'
        
        data = [
            ["Analysis Period", f"{filters.get('start_date', 'N/A')} to {filters.get('end_date', 'N/A')}"],
            ["Departments", departments_str],
            ["Stay Duration Range", f"{min_stay} - {max_stay} days"],
        ]
        
        table = Table(data, colWidths=[1.5*inch, 4*inch])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#1f4788')),
            ('TEXTCOLOR', (0, 0), (0, -1), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('TOPPADDING', (0, 0), (-1, -1), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey),
        ]))
        
        return table
    
    def _create_stay_statistics_table(self, data: Dict) -> Table:
        """Create statistics table for PDF."""
        statistics = data.get('statistics', {}).get('data', {})
        
        table_data = [
            ["Statistic", "Value"],
            ["Total Patients", f"{statistics.get('total_patients', 0):,}"],
            ["Average Stay", f"{statistics.get('average_stay_days', 0):.2f} days"],
            ["Median Stay", f"{statistics.get('median_stay_days', 0):.2f} days"],
            ["Minimum Stay", f"{statistics.get('min_stay_days', 0):.2f} days"],
            ["Maximum Stay", f"{statistics.get('max_stay_days', 0):.2f} days"],
            ["Standard Deviation", f"{statistics.get('std_dev_days', 0):.2f} days"],
        ]
        
        table = Table(table_data, colWidths=[2*inch, 2*inch])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f4788')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('ALIGN', (1, 1), (-1, -1), 'RIGHT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f0f0f0')]),
        ]))
        
        return table
    
    def _create_stay_distribution_table(self, data: Dict) -> Table:
        """Create distribution table for PDF."""
        distribution = data.get('distribution', {}).get('data', [])
        
        if not distribution:
            return None
        
        table_data = [["Duration Range", "Patient Count", "Percentage"]]
        total_patients = sum(item.get('patient_count', 0) for item in distribution)
        
        for item in distribution:
            count = item.get('patient_count', 0)
            pct = (count / total_patients * 100) if total_patients > 0 else 0
            table_data.append([
                item.get('stay_duration_range', 'N/A'),
                str(count),
                f"{pct:.1f}%"
            ])
        
        table = Table(table_data, colWidths=[2*inch, 1.5*inch, 1*inch])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f4788')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('ALIGN', (1, 1), (-1, -1), 'RIGHT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f0f0f0')]),
        ]))
        
        return table
    
    def _create_stay_department_table(self, data: Dict) -> Table:
        """Create department table for PDF."""
        by_department = data.get('by_department', {}).get('data', [])
        
        if not by_department:
            return None
        
        table_data = [["Department", "Avg Stay (days)", "Patients"]]
        
        for dept in sorted(by_department, key=lambda x: x.get('average_stay_days', 0), reverse=True)[:15]:
            table_data.append([
                dept.get('department_name', f"Dept {dept.get('department_id', 'N/A')}"),
                f"{dept.get('average_stay_days', 0):.1f}",
                str(dept.get('patient_count', 0))
            ])
        
        table = Table(table_data, colWidths=[2.5*inch, 1.2*inch, 0.8*inch])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f4788')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('ALIGN', (1, 1), (-1, -1), 'RIGHT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f0f0f0')]),
        ]))
        
        return table
    
    def _create_stay_trends_table(self, data: Dict) -> Table:
        """Create trends table for PDF."""
        trends = data.get('trends', {}).get('monthly', {}).get('data', [])
        
        if not trends:
            return None
        
        table_data = [["Period", "Avg Stay", "Patients", "Change %"]]
        prev_avg = None
        
        for trend in trends[:12]:
            period = trend.get('period', 'N/A')
            avg_stay = trend.get('average_stay_days', 0)
            count = trend.get('patient_count', 0)
            
            change = "—"
            if prev_avg is not None and prev_avg > 0:
                pct_change = ((avg_stay - prev_avg) / prev_avg) * 100
                change = f"{pct_change:+.1f}%"
            
            table_data.append([period, f"{avg_stay:.1f}", str(count), change])
            prev_avg = avg_stay
        
        table = Table(table_data, colWidths=[1.2*inch, 1*inch, 0.8*inch, 0.8*inch])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f4788')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('ALIGN', (1, 1), (-1, -1), 'RIGHT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f0f0f0')]),
        ]))
        
        return table
    
    def _create_patient_stays_table_pdf(self, data: Dict) -> Table:
        """Create patient stays table for PDF."""
        patient_stays = data.get('patient_stays', {}).get('data', [])
        
        if not patient_stays:
            return None
        
        table_data = [["Rank", "Patient ID", "Stay (days)", "Department"]]
        
        for idx, stay in enumerate(patient_stays[:25], 1):  # Top 25 for PDF
            table_data.append([
                str(idx),
                str(stay.get('patient_id', 'N/A')),
                f"{stay.get('stay_days', 0):.1f}",
                stay.get('department_name', f"Dept {stay.get('department_id', 'N/A')}")[:20]
            ])
        
        table = Table(table_data, colWidths=[0.4*inch, 1*inch, 1*inch, 2*inch])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f4788')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('ALIGN', (0, 0), (0, -1), 'CENTER'),
            ('ALIGN', (2, 1), (2, -1), 'RIGHT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f0f0f0')]),
        ]))
        
        return table
    
    def _create_stay_distribution_chart_pdf(self, data: Dict) -> Image:
        """Create distribution chart for PDF."""
        if not MATPLOTLIB_AVAILABLE:
            return None
        
        try:
            distribution = data.get('distribution', {}).get('data', [])
            if not distribution:
                return None
            
            chart_stream = self._create_stay_distribution_chart(distribution)
            
            if chart_stream:
                return Image(chart_stream, width=5*inch, height=3*inch)
            return None
        except Exception as e:
            self.logger.error(f"Error creating distribution chart for PDF: {e}")
            return None
    
    def _create_stay_department_chart_pdf(self, data: Dict) -> Image:
        """Create department chart for PDF."""
        if not MATPLOTLIB_AVAILABLE:
            return None
        
        try:
            by_department = data.get('by_department', {}).get('data', [])
            if not by_department:
                return None
            
            chart_stream = self._create_stay_department_chart(by_department)
            
            if chart_stream:
                return Image(chart_stream, width=5*inch, height=3.5*inch)
            return None
        except Exception as e:
            self.logger.error(f"Error creating department chart for PDF: {e}")
            return None
    
    def _create_stay_trends_chart_pdf(self, data: Dict) -> Image:
        """Create trends chart for PDF."""
        if not MATPLOTLIB_AVAILABLE:
            return None
        
        try:
            trends = data.get('trends', {}).get('monthly', {}).get('data', [])
            if not trends:
                return None
            
            chart_stream = self._create_stay_trends_chart(trends)
            
            if chart_stream:
                return Image(chart_stream, width=5.5*inch, height=3.5*inch)
            return None
        except Exception as e:
            self.logger.error(f"Error creating trends chart for PDF: {e}")
            return None